from __future__ import annotations

from io import BytesIO

from docxtpl import DocxTemplate
from docx import Document
from docx.oxml.ns import qn

from app.core.config import settings


def _set_run_font(run, font_name: str) -> None:
    run.font.name = font_name
    # Ensure Word applies the font for all scripts (ascii/hAnsi/eastAsia/cs).
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:cs"), font_name)


def _normalize_fonts(doc: Document, font_name: str) -> None:
    # Default/known styles
    try:
        normal = doc.styles["Normal"]
        normal.font.name = font_name
    except Exception:
        pass

    # Apply for all styles where possible
    for style in doc.styles:
        try:
            style.font.name = font_name
        except Exception:
            continue

    def visit_paragraphs(paragraphs) -> None:
        for p in paragraphs:
            for run in p.runs:
                try:
                    _set_run_font(run, font_name)
                except Exception:
                    continue

    visit_paragraphs(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                visit_paragraphs(cell.paragraphs)

    for section in doc.sections:
        header = section.header
        footer = section.footer
        visit_paragraphs(getattr(header, "paragraphs", []))
        visit_paragraphs(getattr(footer, "paragraphs", []))


def render_docx(*, template_docx: bytes, context: dict) -> bytes:
    template_io = BytesIO(template_docx)
    doc = DocxTemplate(template_io)
    doc.render(context)
    out = BytesIO()
    doc.save(out)
    rendered = out.getvalue()

    # Force a consistent font for generated documents (DOCX and later PDF conversion).
    # Note: PDF rendering still depends on the font being available in the renderer container.
    rendered_io = BytesIO(rendered)
    document = Document(rendered_io)
    _normalize_fonts(document, settings.document_font_name)
    normalized_out = BytesIO()
    document.save(normalized_out)
    return normalized_out.getvalue()
